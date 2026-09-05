"""Unified multi-format knowledge ingestion service."""

import logging
from typing import Any

from app.services.document_processing.log_parser import LogParser
from app.services.document_processing.pdf_extractor import PDFExtractor
from app.services.document_processing.table_extractor import TableExtractor
from app.services.ingestion.detector import detect_file_type, detect_language
from app.services.llm.document_analysis import DocumentAnalysisService
from app.services.metadata.metadata_extractor import MetadataExtractor
from app.services.ocr.tesseract_ocr import OCRProcessor

logger = logging.getLogger(__name__)


class NormalizedDocument:
    """Represents an ingested knowledge asset broken down into normalized items."""

    def __init__(
        self,
        source_type: str,
        file_name: str,
        language: str,
        machine_model: str | None,
        detected_error_codes: list[str],
        items: list[dict[str, Any]],
        raw_text: str = "",
        metadata: dict[str, Any] | None = None,
    ):
        self.source_type = source_type
        self.file_name = file_name
        self.language = language
        self.machine_model = machine_model
        self.detected_error_codes = detected_error_codes
        self.items = items  # list of {page, section, content, error_codes, metadata}
        self.raw_text = raw_text
        self.metadata = metadata or {}

    def to_dict(self) -> dict[str, Any]:
        return {
            "source_type": self.source_type,
            "file_name": self.file_name,
            "language": self.language,
            "machine_model": self.machine_model,
            "detected_error_codes": self.detected_error_codes,
            "items_count": len(self.items),
            "metadata": self.metadata,
        }


class MultiFormatIngestionService:
    """Processes PDF, images, CSV, logs, docx, and text into normalized documents."""

    def __init__(self) -> None:
        self.pdf_extractor = PDFExtractor()
        self.log_parser = LogParser()
        self.table_extractor = TableExtractor()
        self.ocr = OCRProcessor()
        self.doc_analysis = DocumentAnalysisService()
        self.metadata_extractor = MetadataExtractor()

    def process_file(
        self,
        file_bytes: bytes,
        file_name: str,
        explicit_machine: str | None = None,
    ) -> NormalizedDocument:
        """Ingest any supported file format and return a normalized document."""
        source_type = detect_file_type(file_name, file_bytes)
        logger.info(f"Ingesting file '{file_name}' detected as '{source_type}' ({len(file_bytes)} bytes)")

        if source_type == "pdf":
            return self._process_pdf(file_bytes, file_name, explicit_machine)
        elif source_type == "image":
            return self._process_image(file_bytes, file_name, explicit_machine)
        elif source_type == "csv":
            return self._process_csv(file_bytes, file_name, explicit_machine)
        elif source_type == "log":
            return self._process_log(file_bytes, file_name, explicit_machine)
        elif source_type == "docx":
            return self._process_docx(file_bytes, file_name, explicit_machine)
        else:
            return self._process_text(file_bytes, file_name, explicit_machine)

    def _process_pdf(self, file_bytes: bytes, file_name: str, explicit_machine: str | None) -> NormalizedDocument:
        result = self.pdf_extractor.process_pdf(file_bytes, file_name)
        full_text = " ".join(p["text"] for p in result["pages"])
        lang = detect_language(full_text[:1000])

        meta = self.metadata_extractor.extract(full_text[:5000], file_name)
        machine_model = explicit_machine or meta.get("primary_machine")

        items = []
        all_errors = set(meta.get("error_codes", []))

        for p in result["pages"]:
            page_meta = self.metadata_extractor.extract(p["text"], file_name)
            p_errors = sorted(list(set(page_meta.get("error_codes", []))))
            all_errors.update(p_errors)
            p_machine = explicit_machine or page_meta.get("primary_machine") or machine_model

            blocks = p.get("blocks", [])
            machine_blocks = []
            if blocks and len(blocks) > 1:
                for blk in blocks:
                    blk_clean = " ".join(blk.split())
                    blk_meta = self.metadata_extractor.extract(blk_clean, file_name)
                    b_machines = blk_meta.get("machine_models", [])
                    b_errors = blk_meta.get("error_codes", [])
                    if b_machines or b_errors:
                        b_mach = explicit_machine or blk_meta.get("primary_machine") or p_machine
                        machine_blocks.append({
                            "source_type": "pdf",
                            "file_name": file_name,
                            "machine_model": b_mach,
                            "page": p["page_number"],
                            "section": p["section"],
                            "content": blk_clean,
                            "error_codes": b_errors,
                            "metadata": {
                                "was_ocr": p["was_ocr"],
                                "has_tables": p["has_tables"],
                                "image_count": p["image_count"],
                                "machine_models": b_machines,
                            },
                        })

            if len(machine_blocks) > 1:
                items.extend(machine_blocks)
            else:
                items.append({
                    "source_type": "pdf",
                    "file_name": file_name,
                    "machine_model": p_machine,
                    "page": p["page_number"],
                    "section": p["section"],
                    "content": p["text"],
                    "error_codes": p_errors,
                    "metadata": {
                        "was_ocr": p["was_ocr"],
                        "has_tables": p["has_tables"],
                        "image_count": p["image_count"],
                        "machine_models": page_meta.get("machine_models", []),
                    },
                })

        return NormalizedDocument(
            source_type="pdf",
            file_name=file_name,
            language=lang,
            machine_model=machine_model,
            detected_error_codes=sorted(list(all_errors)),
            items=items,
            raw_text=full_text,
            metadata={
                "total_pages": result["total_pages"],
                "tables_detected": result["tables_detected"],
                "diagrams_detected": result["diagrams_detected"],
                "ocr_pages_processed": result["ocr_pages_processed"],
            },
        )

    def _process_image(self, file_bytes: bytes, file_name: str, explicit_machine: str | None) -> NormalizedDocument:
        # Run OCR first
        ocr_text = self.ocr.extract_text_from_image(file_bytes)
        # Also run vision analysis
        vision = self.doc_analysis.analyze_image_or_screenshot(file_bytes, file_name)

        combined_text = f"{ocr_text}\n{vision.get('extracted_text', '')}\n{vision.get('description', '')}".strip()
        lang = detect_language(combined_text)
        meta = self.metadata_extractor.extract(combined_text, file_name)

        detected_errors = sorted(list(set(vision.get("error_codes", []) + meta.get("error_codes", []))))
        machine_model = explicit_machine or vision.get("machine_model") or meta.get("primary_machine")

        items = [
            {
                "source_type": "image",
                "file_name": file_name,
                "machine_model": machine_model,
                "page": 1,
                "section": vision.get("image_type", "Image Analysis"),
                "content": combined_text or f"Machine image: {file_name}",
                "error_codes": detected_errors,
                "metadata": {
                    "image_type": vision.get("image_type"),
                    "symptoms": vision.get("symptoms", []),
                },
            }
        ]

        return NormalizedDocument(
            source_type="image",
            file_name=file_name,
            language=lang,
            machine_model=machine_model,
            detected_error_codes=detected_errors,
            items=items,
            raw_text=combined_text,
            metadata={
                "image_type": vision.get("image_type"),
                "symptoms": vision.get("symptoms", []),
            },
        )

    def _process_csv(self, file_bytes: bytes, file_name: str, explicit_machine: str | None) -> NormalizedDocument:
        result = self.table_extractor.parse_csv(file_bytes, file_name)
        text = result["content"]
        lang = detect_language(text)
        meta = self.metadata_extractor.extract(text, file_name)
        machine_model = explicit_machine or meta.get("primary_machine")

        items = [
            {
                "source_type": "csv",
                "file_name": file_name,
                "machine_model": machine_model,
                "page": 1,
                "section": "Sensor Data / CSV Table",
                "content": text,
                "error_codes": meta.get("error_codes", []),
                "metadata": {
                    "row_count": result["row_count"],
                    "columns": result["columns"],
                },
            }
        ]

        return NormalizedDocument(
            source_type="csv",
            file_name=file_name,
            language=lang,
            machine_model=machine_model,
            detected_error_codes=meta.get("error_codes", []),
            items=items,
            raw_text=text,
            metadata={"row_count": result["row_count"], "columns": result["columns"]},
        )

    def _process_log(self, file_bytes: bytes, file_name: str, explicit_machine: str | None) -> NormalizedDocument:
        text = file_bytes.decode("utf-8", errors="replace")
        parsed = self.log_parser.parse_log_text(text, file_name)
        lang = detect_language(text[:1000])

        machine_model = explicit_machine or (
            parsed["detected_machines"][0] if parsed["detected_machines"] else None
        )

        items = [
            {
                "source_type": "log",
                "file_name": file_name,
                "machine_model": machine_model,
                "page": 1,
                "section": "Machine Diagnostics Log",
                "content": text,
                "error_codes": parsed["detected_error_codes"],
                "metadata": {
                    "total_log_lines": parsed["total_lines"],
                    "detected_machines": parsed["detected_machines"],
                },
            }
        ]

        return NormalizedDocument(
            source_type="log",
            file_name=file_name,
            language=lang,
            machine_model=machine_model,
            detected_error_codes=parsed["detected_error_codes"],
            items=items,
            raw_text=text,
            metadata={
                "total_log_lines": parsed["total_lines"],
                "detected_machines": parsed["detected_machines"],
            },
        )

    def _process_docx(self, file_bytes: bytes, file_name: str, explicit_machine: str | None) -> NormalizedDocument:
        import io
        import docx
        doc = docx.Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        full_text = "\n\n".join(paragraphs)
        lang = detect_language(full_text[:1000])
        meta = self.metadata_extractor.extract(full_text, file_name)
        machine_model = explicit_machine or meta.get("primary_machine")

        items = [
            {
                "source_type": "docx",
                "file_name": file_name,
                "machine_model": machine_model,
                "page": 1,
                "section": "General Documentation",
                "content": full_text,
                "error_codes": meta.get("error_codes", []),
                "metadata": {"paragraph_count": len(paragraphs)},
            }
        ]

        return NormalizedDocument(
            source_type="docx",
            file_name=file_name,
            language=lang,
            machine_model=machine_model,
            detected_error_codes=meta.get("error_codes", []),
            items=items,
            raw_text=full_text,
            metadata={"paragraph_count": len(paragraphs)},
        )

    def _process_text(self, file_bytes: bytes, file_name: str, explicit_machine: str | None) -> NormalizedDocument:
        text = file_bytes.decode("utf-8", errors="replace")
        lang = detect_language(text[:1000])
        meta = self.metadata_extractor.extract(text, file_name)
        machine_model = explicit_machine or meta.get("primary_machine")

        items = [
            {
                "source_type": "text",
                "file_name": file_name,
                "machine_model": machine_model,
                "page": 1,
                "section": "Technical Notes",
                "content": text,
                "error_codes": meta.get("error_codes", []),
                "metadata": {"length": len(text)},
            }
        ]

        return NormalizedDocument(
            source_type="text",
            file_name=file_name,
            language=lang,
            machine_model=machine_model,
            detected_error_codes=meta.get("error_codes", []),
            items=items,
            raw_text=text,
            metadata={"length": len(text)},
        )
